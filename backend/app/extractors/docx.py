"""
Word document text extractor using python-docx.
"""

from docx import Document


def extract_docx(file_path: str) -> str:
    """
    Extract text from a Word document.
    
    Args:
        file_path: Path to the .docx file
    
    Returns:
        Extracted text content
    """
    doc = Document(file_path)
    text_parts = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(" | ".join(row_text))
    
    return "\n\n".join(text_parts)
